
import base64
from io import BytesIO
import os
import re
import time
import cv2
import json
from flask import Flask, request,jsonify,send_file
import numpy as np
import csv
import numpy as np
import speech_recognition as sr
import random
from docx import Document
from docx.shared import Inches
import base64
import docx2pdf
import firebase_admin
from firebase_admin import credentials
from firebase_admin import firestore
from firebase_admin import storage
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH


cred = credentials.Certificate("serviceAccountKey.json")
firebase_admin.initialize_app(cred,{'storageBucket': 'cyforge-ecaac.appspot.com'})



db=firestore.client()

bucket = storage.bucket()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'    

@app.route('/process-image', methods=['POST'])
def process_image():
    image_data = request.get_json()['image']
    image = cv2.imdecode(np.frombuffer(base64.b64decode(image_data), np.uint8), cv2.IMREAD_UNCHANGED)

    hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
    cv2.imshow("hsv",hsv)
    # define range of blue color in HSV
    lower_white = np.array([0, 0, 231])
    upper_white = np.array([180, 8, 255])
    lower_green = np.array([36, 0, 200])
    upper_green = np.array([89, 155, 255
    ])
    # Create a mask. Threshold the HSV image to get only yellow colors
    mask1 = cv2.inRange(hsv, lower_white, upper_white)
    mask2 = cv2.inRange(hsv, lower_green, upper_green)
    # Bitwise-AND mask and original image
    result_sender = cv2.bitwise_and(image,image, mask= mask1)
    result_receiver = cv2.bitwise_and(image,image, mask= mask2)
    image=result_sender
    image_recv=result_receiver
    # Encode the processed image back to base64
    _, img_encodedrecv = cv2.imencode('.jpg', image_recv)
    _, img_encoded = cv2.imencode('.jpg', image)
    processed_image_data = base64.b64encode(img_encoded.tobytes()).decode('utf-8')
    processed_image_data_recv = base64.b64encode(img_encodedrecv.tobytes()).decode('utf-8')

    return json.dumps({ 'image': processed_image_data,'image_recv':processed_image_data_recv })

@app.route('/convert',methods=['POST'])
def converttotext():
    file = request.files['file']
    suffix = random.randint(100000,999999)
    suffix = str(suffix)
    filename = f'audio_{suffix}.wav'
    file.save(filename)

    try:
        r = sr.Recognizer()
        with sr.AudioFile(filename) as source:
            data = r.record(source)
            text = r.recognize_google(data)

    except:
        return  f"Some error occured"

    res = jsonify(text=text)
    return res

@app.route('/receive_data', methods=['POST'])
def receive_data():
    # read the chatProcessed data from the request
    chat_processed = request.form.get('chatProcessed', '')

    # read the recognizedText data from the request
    recognized_text = request.form.get('recognizedText', '')

    # read the controller_sender data from the request
    controller_sender = request.form.get('controller_sender', '')

    # read the controller_receiver data from the request
    controller_receiver = request.form.get('controller_receiver', '')
    filename=request.form.get('filename','')
    email=request.form.get('email','')
    print(email)

    # create a new Word document
    document =  Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Cambria'
    style_bold = document.styles.add_style('Bold', WD_STYLE_TYPE.PARAGRAPH)
    style_bold.font.bold = True
    style_bold.font.size = Pt(14)

    style_underline = document.styles.add_style('Underline', WD_STYLE_TYPE.CHARACTER)
    style_underline.font.underline = True

    style_large = document.styles.add_style('MyStyleLarge', WD_STYLE_TYPE.PARAGRAPH)
    style_large.font.size = Pt(26)

    style_small = document.styles.add_style('MyStyleSmall', WD_STYLE_TYPE.PARAGRAPH)

    # exec_style = document.styles.add_style('ExecStyle',1)
    # exec_style.font.color.rgb = RGBColor(255,0,0)
    # exec_style.font.size = Pt(14)
    # exec_style.paragraph_format.background_color.rgb = RGBColor(255,255,0)

    # Add centered heading to the document
    heading = document.add_heading('CYBER FORENSIC ANALYSIS REPORT', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.style = 'MyStyleLarge'
    heading.runs[0].underline = True
    heading.runs[0].bold = True

    # Add centered paragraph to the document
    paragraph1 = document.add_paragraph('Forensic Report u/s Indian Evidence Act Sec "65" (B)')
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph1.style = 'Bold'
    paragraph1.style.font.size = Pt(14)
    paragraph1.runs[0].bold = True

    # Add left aligned paragraph to the document
    paragraph2 = document.add_paragraph('\nSubject: ')
    paragraph2.style = 'Bold'
    paragraph2.add_run('\n\n')


    paragraph3 = document.add_paragraph(f"1. Submitted by: \n\n2. Experience: \n\n\nTools/Instruments/Equipments: Cyforge, \n\nSupporting Details: \n\nExamination: \n\n")
    paragraph3.style = style
    paragraph3.style.font.size = Pt(14)

    document.add_page_break()
    paragraph4 = document.add_paragraph('Executive Summary')
    document.add_page_break()

    paragraph5 = document.add_paragraph('\n1. Analyzation: \n\nIn forensic science, audio-video forensics forms three basic principles such as acquisition, analysis, and evaluation of audio and video recordings which are admissible in the court of law. One of the main tasks of audio and video forensic experts is to establish the authenticity and credibility of digital evidence.')
    paragraph5.add_run("\n\n What is hash value? \n\nA hash value is a numeric value of a fixed length that uniquely identifies data. Hash values represent large amounts of data as much smaller numeric values, so they are used with digital signatures. You can sign a hash value more efficiently than signing the larger value.")
    paragraph5.add_run("\n\n How is hash value calculated? \n\n Hashing is simply passing some data through a formula that produces a result, called a hash. That hash is usually a string of characters and the hashes generated by a formula are always the same length, regardless of how much data you feed into it. For example, the MD5 formula always produces 32 character-long hashes.")


    ### METADATA
    # metadata = document.add_paragraph()

    # for i,file in enumerate(audio_files):
    #     size = os.path.getsize(file)
    #     last_modified = get_date(os.path.getmtime(file))
    #     create_date = get_date(os.path.getctime(file))
    #     sha256sum , md5sum = get_hashes(file)
    #     filename = file.split('/')[-1]

    #     with wave.open(file,'rb') as audio_file:
    #         num_frames = audio_file.getnframes()
    #         frame_rate = audio_file.getframerate()
    #         duration = num_frames / float(frame_rate)


    #     print(f"\n\n\n\n {i}) {filename} \t metadata \n\n size ---> {size} bytes \n last modified ---> {last_modified} \n creation date ---> {create_date}")
    #     print(" duration --> {:.2f}s".format(duration) + f"\n frame rate ---> {frame_rate}")
    #     print(f" sha256sum ---> {sha256sum} \n md5sum ---> {md5sum}\n______________________________")

    #     metadata.add_run(f" \n\n\n {i+1}) {filename} metadata \n size ---> {size} bytes \n last modified ---> {last_modified} \n creation date ---> {create_date}\n")
    #     metadata.add_run(" duration --> {:.2f}s".format(duration) + f"\n frame rate ---> {frame_rate} \n")
    #     metadata.add_run(f" sha256sum ---> {sha256sum} \n md5sum ---> {md5sum}\n___________________________________________________________________\n\n\n")


    document.add_page_break()

    paragraph6 = document.add_paragraph("\nReporting: \n\nThe results of the analysis should be reported. Items to be reported may include the following: a description of the actions employed; an explanation of how tools and procedures were selected; a determination of any other actions that should be performed, such as forensic examination of additional data sources, securing identified vulnerabilities, and improving existing security controls; and recommendations for improvements to policies, guidelines, procedures, tools, and other aspects of the forensic process.\
    \n\nDigital and multimedia evidence, as well as the tools, techniques, and methodologies used in an examination, are subject to challenge in a court of law or other formal proceedings. Proper documentation is essential in providing individuals the ability to reproduce the forensic process and the results. The Scientific Working Group on Digital Evidence (SWGDE) states that Reporting is the process of preparing a summary of steps taken during the examination of digital media. A thorough examination report is written using documentation collected by the examiner, including photographs, drawings, case-notes, tool-generated content, etc. Many forensic tools come with built-in reporting functionality that is specific to that tool’s actions and results, but does not typically document the full scope of the examination. Tool reports may be considered supporting documentation to the examination report or referenced as an appendix.\
    \n\nThe purpose of the Forensic Report is simply to tell the story of what the presence or absence of the digital artifact indicates, regardless, if it is inculpatory or exculpatory in nature.")
    document.add_page_break()

    paragraph7 = document.add_paragraph("\n \t\t\t\t   Verification ")
    paragraph7.style = 'Bold'
    paragraph7.add_run("\n\n\n\n\n")

    paragraph8 = document.add_paragraph("\n Conclusion: ")
    paragraph8.style = 'Bold'
    paragraph8.add_run("\n\n")
    document.add_page_break()

    paragraph9 = document.add_paragraph("\n EXPERT WITNESS TESTIMONY \n\n")
    paragraph9.add_run("What is a Video Forensic Expert Witness?\n \
    A Video Forensic Expert has the scientific knowledge, training and expertise necessary to enhance and authenticate video recordings that are being used in a criminal or civil court case. The expert witness has previous court room experience testifying and helps the Trier of Fact understand the video evidence that is offered in a litigation.")
    paragraph9.add_run("\n\n\n AUDIO ANALYSIS / AUTHENTICATION \n\n")
    paragraph9.add_run("The first step of an analysis is for the examiner to simply listen to or view the recorded footage. The examiner will then begin to locate the area of interest to be enhanced and examined in closer detail using specialized devices and software.\
    \n\nBefore processing audio and video evidence, a working copy of the evidence may be created. This assures that the original evidence is always available in its unaltered state. In addition, the original will always be available for comparison to the processed copy.\
    \n\nAll examination procedures are carefully constructed so that the image or video is a true and accurate representation of the scene. Investigators never change the recorded data—they only enhance what is already present.\
    \n\nThe primary aspects of audio forensics are establishing the authenticity of audio evidence, performing enhancement of audio recordings to improve speech intelligibility and the audibility of low-level sounds, and interpreting and documenting sonic evidence, such as identifying talkers, transcribing dialog, and reconstructing crime or accident scenes and timelines.\
    \n\nModern audio forensics makes extensive use of digital signal processing, with the former use of analog filters now being obsolete. Techniques such as adaptive filtering and discrete Fourier transforms are used extensive.\
    \n\nFrequency Equalization - Highly precise equalizers can be used to boost or cut specific bands of frequencies. To help make speech more intelligible, the frequency band containing most speech content, 200Hz–5000Hz, can be amplified or isolated. If amplification is applied to a frequency range, other information residing in this frequency range will be boosted as well. If noise resides in this same range, this noise will also be increased, limiting the ability to clarify voices.\
    \n\nLoud background noises may be analyzed by a spectrum analyzer and the corresponding frequencies reduced so that these noises are less noticeable.\
    \n\nCompression - Faint sounds in the recording can be boosted by compressing or leveling the signal so that the dynamic range of the material is reduced, making soft sounds more apparent.")


    image_files = request.files.getlist('imageFiles[]')
    for i, image_file in enumerate(image_files):
        # save the image file to disk
        image_file.save(f'image_{i}.jpg')
        
        # add the image to the document
        # document.add_picture(f'image_{i}.jpg', width=Inches(4))
    # add the chatProcessed data to the document
    if chat_processed:
        document.add_paragraph(chat_processed)
    else:
        document.add_paragraph("The chatProcessed variable is not received.")

    # add the recognizedText data to the document
    if recognized_text:
        document.add_paragraph(recognized_text)
    else:
        document.add_paragraph("The recognizedText variable is not received.")

    # add the controller_sender data to the document
    if controller_sender:
        document.add_paragraph(controller_sender)
    else:
        document.add_paragraph("The controller_sender variable is not received.")

    # add the controller_receiver data to the document
    if controller_receiver:
        document.add_paragraph(controller_receiver)
    else:
        document.add_paragraph("The controller_receiver variable is not received.")
    


    # save the document to disk
    document.save(f'{filename}.docx')
    doc_file = f'{filename}.docx'
 
   
    blob = bucket.blob(doc_file)
    blob.upload_from_filename(doc_file)
    blob.make_public()
    print("your file url", blob.public_url)
    data = {"fileName": os.path.basename(doc_file), "fileUrl": blob.public_url}
    db.collection(email).document().set(data)   
           
          
    
    return "send_file( doc_file,as_attachment=True)"


@app.route('/text',methods=['POST'])
def process_txt_file():
    # Get the uploaded file from the request
    file = request.files['file']
    
    # Read the contents of the file
    contents = file.read().decode('utf-8')
    
    # Split the contents into individual lines
    lines = contents.split('\n')
    
    # Initialize a dictionary to store the data for each sender
    senders = {}
    
    # Loop through each line and split it into the timestamp, sender name, and message text
    for line in lines:
        # Use a regular expression to match the timestamp and sender name
        match = re.match(r'(\d{2}/\d{2}/\d{4}, \d{1,2}:\d{2}\s*[ap]m) - (.+?): (.+)', line)
        if match:
            timestamp = match.group(1)
            sender_name = match.group(2)
            message_text = match.group(3)
            
            # Check if the sender name already exists in the dictionary
            if sender_name in senders:
                # If it does, add the message to the sender's list of messages
                senders[sender_name].append({'timestamp': timestamp, 'message_text': message_text})
            else:
                # If it doesn't, create a new list for the sender and add the message to it
                senders[sender_name] = [{'timestamp': timestamp, 'message_text': message_text}]
    
    # Return the data for each sender as a JSON response
    return jsonify(senders)


if __name__ == '__main__':
    app.run(host='0.0.0.0')

